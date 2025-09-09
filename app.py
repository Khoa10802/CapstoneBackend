import io
import os
import time  # 🆕 Thêm để đo thời gian

from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from docx import Document
from solcx.exceptions import SolcError, UnsupportedVersionError

# Custom imports
from sol_compilation import package_assemble, ExternalInclusionError, VersionNotFoundError
from predict import predict_vulnerabilities, predict_contract  # Gộp cả 2 hàm predict

app = Flask(__name__, template_folder="templates")
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ─────────────────────────────────────────────
# GET route to serve the main page
@app.route("/", methods=["GET"])
def index():
    return render_template('index.html')

# ─────────────────────────────────────────────
# POST route to upload file and run vulnerability prediction
@app.route("/upload", methods=["POST"])
def upload():
    if 'file' not in request.files:
        return jsonify({"error": "File not found"}), 400

    sol_file = request.files['file']
    if sol_file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    try:
        start_time = time.time()  # ⏱️ Bắt đầu tính thời gian

        # ✅ Always save file first, so rescan can use it later
        file_path = os.path.join(UPLOAD_FOLDER, sol_file.filename)
        sol_file.save(file_path)

        if sol_file.filename.lower().endswith('.docx'):
            # Read and extract text from docx
            doc = Document(file_path)
            docx_text = [para.text for para in doc.paragraphs]
            if len(docx_text) <= 1 and docx_text[0] == '':
                return jsonify({"error": f"{sol_file.filename} is maybe empty"}), 400
            compiled_sol = package_assemble('\n'.join(docx_text))
        else:
            # Handle raw Solidity (.sol) files
            with open(file_path, "r", encoding="utf-8") as f:
                sol_str = f.read()
            if len(sol_str) <= 0:
                return jsonify({"error": f"{sol_file.filename} is maybe empty"}), 400
            compiled_sol = package_assemble(sol_str)

        # Run ML prediction
        result = predict_vulnerabilities(compiled_sol)

        end_time = time.time()
        duration_seconds = round(end_time - start_time, 2)  # ⏱️ Tính thời gian chạy

        # ✅ Add metadata for frontend
        result["duration"] = duration_seconds
        result["filename"] = sol_file.filename

        return result

    # Handle custom errors
    except ExternalInclusionError:
        return jsonify({"error": f"{sol_file.filename} may contain external library"}), 400
    except VersionNotFoundError:
        return jsonify({"error": f"{sol_file.filename} may not define pragma version"}), 400
    except SolcError:
        return jsonify({"error": "Compilation failed"}), 400
    except UnsupportedVersionError:
        return jsonify({"error": "Does not support this file version"}), 400
    except ValueError:
        return jsonify({"error": "No contracts found"}), 400

# ─────────────────────────────────────────────
# POST route to re-run prediction on an already uploaded file
@app.route("/rescan", methods=["POST"])
def rescan():
    data = request.get_json()
    if not data or "filename" not in data:
        return jsonify({"error": "Filename required"}), 400

    filename = data["filename"]
    file_path = os.path.join(UPLOAD_FOLDER, filename)

    if not os.path.exists(file_path):
        return jsonify({"error": f"File {filename} not found"}), 404

    try:
        start_time = time.time()

        if filename.lower().endswith('.docx'):
            # Reload from saved docx
            doc = Document(file_path)
            docx_text = [para.text for para in doc.paragraphs]
            if len(docx_text) <= 1 and docx_text[0] == '':
                return jsonify({"error": f"{filename} is maybe empty"}), 400
            compiled_sol = package_assemble('\n'.join(docx_text))
        else:
            # Reload from saved .sol file
            with open(file_path, "r", encoding="utf-8") as f:
                sol_str = f.read()
            if len(sol_str) <= 0:
                return jsonify({"error": f"{filename} is maybe empty"}), 400
            compiled_sol = package_assemble(sol_str)

        # Run ML prediction
        result = predict_vulnerabilities(compiled_sol)

        end_time = time.time()
        duration_seconds = round(end_time - start_time, 2)

        # ✅ Add metadata again for consistency
        result["duration"] = duration_seconds
        result["filename"] = filename

        return result

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ─────────────────────────────────────────────
# Optional: Raw Solidity string (via API)
@app.route('/predict', methods=['POST'])
def predict():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Empty filename'}), 400

    try:
        source_code = file.read().decode('utf-8')
        result = predict_contract(source_code)
        return jsonify({'result': result})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ─────────────────────────────────────────────
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000)
